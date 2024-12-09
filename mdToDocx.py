from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import markdown
import re

def set_rtl_list_style(doc, numbered=False):
    # Create RTL list style name based on type
    style_name = 'RTL Number List' if numbered else 'RTL List'
    
    # Create RTL list style
    list_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    list_style.base_style = doc.styles['Normal']
    list_style.font.rtl = True
    list_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    list_style.paragraph_format.left_indent = Inches(0.25)
    list_style.paragraph_format.first_line_indent = Inches(-0.25)
    return list_style

def set_paragraph_rtl(paragraph):
    """Set RTL for paragraph and its runs"""
    # Set paragraph direction RTL
    pPr = paragraph._element.get_or_add_pPr()
    pPr.set(qn('w:bidi'), '1')
    
    # Set run direction RTL for all runs in the paragraph
    for run in paragraph.runs:
        rPr = run._element.get_or_add_rPr()
        rPr.set(qn('w:rtl'), '1')

def process_bold_text(text):
    """Process text and return list of (text, is_bold) tuples"""
    parts = []
    pattern = r'\*\*(.*?)\*\*'
    last_end = 0
    
    for match in re.finditer(pattern, text):
        # Add non-bold text before the match
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False))
        # Add bold text
        parts.append((match.group(1), True))
        last_end = match.end()
    
    # Add remaining non-bold text
    if last_end < len(text):
        parts.append((text[last_end:], False))
    
    return parts if parts else [(text, False)]

def create_rtl_document(markdown_content):
    # Create a new document
    doc = Document()
    
    # Convert markdown to HTML
    html_content = markdown.markdown(markdown_content, extensions=['tables'])
    
    # Split content by headers
    sections = re.split(r'(#+\s.*)', markdown_content)
    
    # Set up RTL paragraph style
    style = doc.styles.add_style('RTL', WD_STYLE_TYPE.PARAGRAPH)
    style.font.rtl = True
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Set up RTL list styles
    list_style = set_rtl_list_style(doc, numbered=False)
    numbered_list_style = set_rtl_list_style(doc, numbered=True)
    
    # Process each section
    for section in sections:
        if not section.strip():
            continue
            
        # Check if it's a header
        if section.strip().startswith('#'):
            level = len(re.match(r'^#+', section.strip()).group())
            text = section.strip('#').strip()
            
            # Add header with appropriate style
            p = doc.add_paragraph(style='Heading {}'.format(min(level, 9)))
            run = p.add_run(text)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            set_paragraph_rtl(p)
            
        else:
            # Process regular paragraphs
            paragraphs = section.split('\n')
            list_number = 1  # Counter for numbered lists
            
            for para in paragraphs:
                if para.strip():
                    # Check if it's a numbered list item
                    if re.match(r'^\d+\.', para.strip()):
                        p = doc.add_paragraph(style='RTL Number List')
                        text = re.sub(r'^\d+\.\s*', '', para.strip())
                        # Process bold text in list item
                        for part_text, is_bold in process_bold_text(text):
                            run = p.add_run(part_text + ' ')
                            run.bold = is_bold
                        p.add_run(f'.{list_number}')  # Add number at the end for RTL
                        set_paragraph_rtl(p)
                        list_number += 1
                    # Check if it's a bullet list item
                    elif para.strip().startswith('- '):
                        p = doc.add_paragraph(style='RTL List')
                        text = para.strip('- ').strip()
                        # Process bold text in list item
                        for part_text, is_bold in process_bold_text(text):
                            run = p.add_run(part_text + ' ')
                            run.bold = is_bold
                        p.add_run('•')  # Add bullet at the end for RTL
                        set_paragraph_rtl(p)
                    else:
                        p = doc.add_paragraph(style='RTL')
                        # Process bold text in paragraph
                        for part_text, is_bold in process_bold_text(para.strip()):
                            run = p.add_run(part_text)
                            run.bold = is_bold
                        set_paragraph_rtl(p)
    
    return doc

def main():
    # Read the markdown file
    with open('tourism.md', 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Create document
    doc = create_rtl_document(markdown_content)
    
    # Save the document
    doc.save('tourism_rtl.docx')
    print("Document has been created successfully!")

if __name__ == "__main__":
    main()
